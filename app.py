import pandas as pd
import streamlit as st
from difflib import SequenceMatcher
from io import BytesIO
from docx import Document

@st.cache(show_spinner=False)
def get_similarity(a, b):
    return SequenceMatcher(None, a, b).ratio()

@st.cache(show_spinner=False)
def load_xlsx_bilan(xlsx_file):
    dc_matieres = pd.read_excel(xlsx_file,None,skiprows=2)

    liste_eleves = []
    
    ################################
    # lecture bulletins par matière
    ################################
    
    for matiere in dc_matieres.keys():
        
        df_matiere = dc_matieres[matiere]
    
        df_matiere = df_matiere.dropna(how='all',axis=0)
        df_matiere = df_matiere.dropna(how='all',axis=1)
        df_matiere['NomPrenom'] = df_matiere['Unnamed: 0'] + ' ' + df_matiere['Unnamed: 1']
        df_matiere = df_matiere.set_index('NomPrenom')
        df_matiere = df_matiere.drop(columns=['Unnamed: 0','Unnamed: 1'])
        
        df_matiere = df_matiere.dropna(how='all',axis=0)
        df_matiere = df_matiere.dropna(how='all',axis=1)
        
        curr_list_eleves = list(df_matiere.index)
        
        for eleve in curr_list_eleves:
            if not pd.isna(eleve) and eleve not in liste_eleves:
                liste_eleves.append(eleve)
        
        dc_matieres[matiere] = df_matiere
    
    return dc_matieres, liste_eleves

@st.cache(show_spinner=False)
def gen_bulletin(dc_matieres,liste_eleves):
    ##################################
    # construction bulletin par élève
    ##################################
    dc_bulletin = {}
    
    for eleve in liste_eleves:
        dc_bulletin[eleve] = {}
        for matiere in dc_matieres.keys():
            dc_bulletin[eleve][matiere] = {}
            df_matiere = dc_matieres[matiere]
            
            if eleve in df_matiere.index:
                rubriques = df_matiere.columns
                for rubrique in rubriques:
                    appreciation = df_matiere.loc[eleve,rubrique]
                    if not pd.isna(appreciation):
                        dc_bulletin[eleve][matiere][rubrique] = appreciation
    
    # purge des matières vides
    for eleve in dc_bulletin.keys():
        for matiere in dc_matieres.keys():
            if len(dc_bulletin[eleve][matiere].keys()) < 1:
                del dc_bulletin[eleve][matiere]
    
    # vérification des noms d'élèves similaires:
    new_list = []
    for eleve in liste_eleves:
        pas_doublon = True
        for ref_eleve in new_list:
            similarity = get_similarity(eleve,ref_eleve)
            if similarity >= 0.8:
                dc_bulletin[ref_eleve] = {**dc_bulletin[ref_eleve], **dc_bulletin[eleve]}
                del dc_bulletin[eleve]
                pas_doublon = False
        if(pas_doublon):
            new_list.append(eleve)
    
    return dc_bulletin


def display_bulletin_eleve(dc_eleve):
    for matiere in dc_eleve.keys():
        st.subheader(matiere)
        for rubrique in dc_eleve[matiere].keys():
            st.markdown('+ **'+rubrique+' : **' + dc_eleve[matiere][rubrique])

def generate_docx_bulletin_eleve(doc,eleve,dc_eleve):
    doc.add_heading('Rapport '+eleve, 0)
    for matiere in dc_eleve.keys():
        doc.add_heading(matiere, level=1)
        for rubrique in dc_eleve[matiere].keys():
            p = doc.add_paragraph('', style='List Bullet')
            p.add_run(rubrique+' : ').bold = True
            p.add_run(dc_eleve[matiere][rubrique])
    
    return doc

def generate_docx(eleve,data):
    output = BytesIO()
    document = Document()
    if eleve:
        document = generate_docx_bulletin_eleve(document,eleve,data)
    else:
        for curr_eleve in data.keys():
            document = generate_docx_bulletin_eleve(document,curr_eleve,data[curr_eleve])
            document.add_page_break()
    document.save(output)
    return output.getvalue()
    



#######################################################################################
    
st.sidebar.title('Génération bulletins élèves')
uploaded_file = st.sidebar.file_uploader("Choix du fichier bilan", type=["xlsx","ods"])

if uploaded_file:
    with st.spinner('Chargement du fichier...'):
        dc_bilan_par_matieres, liste_des_eleves = load_xlsx_bilan(uploaded_file)
    with st.spinner('Préparation des bulletins...'):
        dc_bulletin_tous_eleves = gen_bulletin(dc_bilan_par_matieres, liste_des_eleves)
    
    dwnld_rapport_classe = st.sidebar.download_button(label='Télécharger le rapport de la classe',
                                       data=generate_docx(eleve=False,data=dc_bulletin_tous_eleves),
                                       file_name='rapport_classe.docx',
                                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    
    eleve = st.selectbox('Sélectionner un élève', dc_bulletin_tous_eleves.keys()) 
    
    display_bulletin_eleve(dc_bulletin_tous_eleves[eleve])
    
    dwnld_rapport_eleve = st.download_button(label="Télécharger le rapport de l'élève",
                                       data=generate_docx(eleve=eleve,data=dc_bulletin_tous_eleves[eleve]),
                                       file_name='rapport_'+eleve+'.docx',
                                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
